from __future__ import annotations

import re
from pathlib import Path

from docx import Document


class DOCXFiller:
    def __init__(self, template: str | None = None) -> None:
        self.template = template

    def fill(self, form_data: dict[str, str], output_path: str) -> str:
        output = Path(output_path)
        output.parent.mkdir(parents=True, exist_ok=True)

        if self.template and Path(self.template).exists():
            document = Document(self.template)
            self._replace_placeholders(document, form_data)
        else:
            document = Document()
            document.add_heading("Home Loan Application (Extracted Data)", level=1)
            table = document.add_table(rows=1, cols=2)
            table.rows[0].cells[0].text = "Field"
            table.rows[0].cells[1].text = "Value"
            for key in sorted(form_data.keys()):
                row = table.add_row().cells
                row[0].text = key
                row[1].text = str(form_data[key])

        document.save(str(output))
        return str(output)

    @staticmethod
    def _replace_placeholders(document: Document, form_data: dict[str, str]) -> None:
        DOCXFiller._replace_in_container(document, form_data)
        for section in document.sections:
            DOCXFiller._replace_in_container(section.header, form_data)
            DOCXFiller._replace_in_container(section.footer, form_data)

    @staticmethod
    def _replace_in_container(container, form_data: dict[str, str]) -> None:
        for paragraph in container.paragraphs:
            DOCXFiller._replace_in_paragraph(paragraph, form_data)

        for table in container.tables:
            for row in table.rows:
                for cell in row.cells:
                    DOCXFiller._replace_in_container(cell, form_data)

    @staticmethod
    def _replace_in_paragraph(paragraph, form_data: dict[str, str]) -> None:
        original_text = "".join(run.text for run in paragraph.runs) if paragraph.runs else paragraph.text
        updated_text = DOCXFiller._replace_placeholders_in_text(original_text, form_data)
        if updated_text == original_text:
            return

        if paragraph.runs:
            paragraph.runs[0].text = updated_text
            for run in paragraph.runs[1:]:
                run.text = ""
            return

        paragraph.text = updated_text

    @staticmethod
    def _replace_placeholders_in_text(text: str, form_data: dict[str, str]) -> str:
        updated = text
        for key, value in form_data.items():
            placeholder = re.compile(r"\{\{\s*" + re.escape(key) + r"\s*\}\}")
            updated = placeholder.sub(str(value), updated)
        return updated

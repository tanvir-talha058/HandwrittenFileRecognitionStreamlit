from __future__ import annotations

import os
from pathlib import Path
from typing import Any

import cv2
import fitz
import numpy as np
from docx import Document

from .preprocessor import preprocess_image


class OCREngine:
    def __init__(
        self,
        lang: str = "en",
        use_gpu: bool = False,
        use_angle_cls: bool = True,
        det_db_thresh: float = 0.3,
    ) -> None:
        self.lang = lang
        self.use_gpu = use_gpu
        self.use_angle_cls = use_angle_cls
        self.det_db_thresh = det_db_thresh
        self.last_error: str | None = None
        self._ocr: Any | None = None
        self._force_safe_cpu_runtime = False

    def _get_ocr(self) -> Any:
        if self._ocr is None:
            self._ocr = self._build_ocr()
        return self._ocr

    def _configure_runtime_environment(self) -> None:
        os.environ.setdefault("PADDLE_PDX_DISABLE_MODEL_SOURCE_CHECK", "True")

        if not self.use_gpu or self._force_safe_cpu_runtime:
            os.environ["FLAGS_use_mkldnn"] = "0"
            os.environ["FLAGS_enable_pir_api"] = "0"
            os.environ["FLAGS_enable_pir_in_executor"] = "0"

            try:
                import paddle

                paddle.set_flags(
                    {
                        "FLAGS_use_mkldnn": False,
                        "FLAGS_enable_pir_api": False,
                    }
                )
            except Exception:
                # Flag support varies across Paddle builds; env vars still help.
                pass

    def _build_ocr(self) -> Any:
        self._configure_runtime_environment()

        try:
            from paddleocr import PaddleOCR
        except ImportError as exc:
            raise RuntimeError(
                "PaddleOCR is not installed. Install dependencies with: pip install -r requirements.txt"
            ) from exc

        try:
            return PaddleOCR(
                use_angle_cls=self.use_angle_cls,
                lang=self.lang,
                use_gpu=self.use_gpu,
                det_db_thresh=self.det_db_thresh,
            )
        except ValueError as exc:
            # PaddleOCR >= 3.x removed several legacy constructor args.
            if "Unknown argument" not in str(exc):
                raise

            device = "gpu" if self.use_gpu else "cpu"
            init_kwargs: dict[str, Any] = {
                "lang": self.lang,
                "device": device,
                "use_doc_orientation_classify": False,
                "use_doc_unwarping": False,
                "use_textline_orientation": False,
                "text_det_thresh": self.det_db_thresh,
            }
            if device == "cpu":
                # PaddleOCR 3.x enables MKL-DNN on CPU by default, which crashes in
                # some Paddle runtime combinations with oneDNN executor errors.
                init_kwargs["enable_mkldnn"] = False
            return PaddleOCR(**init_kwargs)

    @staticmethod
    def _is_problematic_runtime_error(exc: Exception) -> bool:
        message = str(exc)
        return any(
            fragment in message
            for fragment in (
                "ConvertPirAttribute2RuntimeAttribute",
                "onednn_instruction.cc",
                "oneDNN",
                "mkldnn",
            )
        )

    def _retry_with_safe_cpu_runtime(self, image: np.ndarray) -> Any:
        self._force_safe_cpu_runtime = True
        self.use_gpu = False
        self._ocr = None
        self._configure_runtime_environment()
        return self._predict(image)

    def _load_first_image(self, input_path: str) -> np.ndarray:
        path = Path(input_path)
        if not path.exists():
            raise FileNotFoundError(f"Input file not found: {input_path}")

        suffix = path.suffix.lower()
        if suffix == ".pdf":
            pages = self._pdf_pages_to_images(path)
            if not pages:
                raise ValueError("PDF has no pages.")
            return pages[0]

        image = cv2.imread(str(path))
        if image is None:
            raise ValueError(f"Unsupported image or unreadable file: {input_path}")
        return image

    @staticmethod
    def _pdf_pages_to_images(path: Path) -> list[np.ndarray]:
        document = fitz.open(path)
        if len(document) == 0:
            document.close()
            raise ValueError("PDF has no pages.")

        images: list[np.ndarray] = []
        for page in document:
            pixmap = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
            image = np.frombuffer(pixmap.samples, dtype=np.uint8).reshape(pixmap.height, pixmap.width, pixmap.n)

            if pixmap.n == 4:
                images.append(cv2.cvtColor(image, cv2.COLOR_RGBA2BGR))
            else:
                images.append(cv2.cvtColor(image, cv2.COLOR_RGB2BGR))
        document.close()
        return images

    @staticmethod
    def _extract_pdf_text_lines(path: Path) -> list[str]:
        lines: list[str] = []
        document = fitz.open(path)
        try:
            for page in document:
                page_text = page.get_text("text")
                if not page_text:
                    continue
                for line in page_text.splitlines():
                    compact = line.strip()
                    if compact:
                        lines.append(compact)
        finally:
            document.close()
        return lines

    @staticmethod
    def _extract_docx_text_lines(path: Path) -> list[str]:
        lines: list[str] = []
        document = Document(str(path))

        for paragraph in document.paragraphs:
            compact = paragraph.text.strip()
            if compact:
                lines.append(compact)

        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text and cell.text.strip()]
                if cells:
                    lines.append(" | ".join(cells))

        return lines

    @staticmethod
    def _extract_txt_lines(path: Path) -> list[str]:
        text = path.read_text(encoding="utf-8", errors="ignore")
        return [line.strip() for line in text.splitlines() if line.strip()]

    @staticmethod
    def _lines_to_blocks(lines: list[str]) -> list[dict[str, Any]]:
        blocks: list[dict[str, Any]] = []
        y = 20
        for line in lines:
            blocks.append(
                {
                    "text": line,
                    "confidence": 1.0,
                    "bbox": [[10, y], [1200, y], [1200, y + 24], [10, y + 24]],
                }
            )
            y += 34
        return blocks

    def run(self, input_path: str, min_confidence: float = 0.8, preprocess: bool = False) -> list[dict[str, Any]]:
        self.last_error = None
        path = Path(input_path)
        suffix = path.suffix.lower()

        if suffix == ".docx":
            return self._lines_to_blocks(self._extract_docx_text_lines(path))
        if suffix == ".txt":
            return self._lines_to_blocks(self._extract_txt_lines(path))

        images: list[np.ndarray]
        if suffix == ".pdf":
            images = self._pdf_pages_to_images(path)
        else:
            images = [self._load_first_image(input_path)]

        try:
            self._get_ocr()
        except Exception as exc:
            self.last_error = f"{type(exc).__name__}: {exc}"
            if suffix == ".pdf":
                return self._lines_to_blocks(self._extract_pdf_text_lines(path))
            return []

        all_blocks: list[dict[str, Any]] = []
        page_stride = 10000

        for page_index, image in enumerate(images):
            page_image = preprocess_image(image) if preprocess else image

            try:
                result = self._predict(page_image)
            except TypeError as exc:
                # PaddleOCR >= 3.x no longer accepts the cls keyword here.
                if "unexpected keyword argument 'cls'" not in str(exc):
                    raise
                try:
                    result = self._ocr.ocr(page_image)
                except Exception as nested_exc:
                    self.last_error = f"{type(nested_exc).__name__}: {nested_exc}"
                    return []
            except Exception as exc:
                if self._is_problematic_runtime_error(exc):
                    try:
                        result = self._retry_with_safe_cpu_runtime(page_image)
                    except Exception as retry_exc:
                        self.last_error = f"{type(retry_exc).__name__}: {retry_exc}"
                        return []
                else:
                    self.last_error = f"{type(exc).__name__}: {exc}"
                    return []

            parsed = self._parse_result(result)
            if not parsed:
                continue

            for item in parsed:
                bbox, payload = item
                text, confidence = payload
                if confidence < min_confidence:
                    continue
                adjusted_bbox = [[int(p[0]), int(p[1]) + (page_index * page_stride)] for p in bbox]
                all_blocks.append(
                    {
                        "text": text.strip(),
                        "confidence": float(confidence),
                        "bbox": adjusted_bbox,
                        "page": page_index + 1,
                    }
                )

        if all_blocks:
            return all_blocks

        # Fallback for digitally generated PDFs where OCR can fail or be unnecessary.
        if suffix == ".pdf":
            return self._lines_to_blocks(self._extract_pdf_text_lines(path))

        return []

    def run_image_array(
        self,
        image: np.ndarray,
        min_confidence: float = 0.8,
        preprocess: bool = False,
        page_number: int = 1,
    ) -> list[dict[str, Any]]:
        self.last_error = None
        page_image = preprocess_image(image) if preprocess else image

        try:
            self._get_ocr()
        except Exception as exc:
            self.last_error = f"{type(exc).__name__}: {exc}"
            return []

        try:
            result = self._predict(page_image)
        except TypeError as exc:
            if "unexpected keyword argument 'cls'" not in str(exc):
                raise
            try:
                result = self._ocr.ocr(page_image)
            except Exception as nested_exc:
                self.last_error = f"{type(nested_exc).__name__}: {nested_exc}"
                return []
        except Exception as exc:
            if self._is_problematic_runtime_error(exc):
                try:
                    result = self._retry_with_safe_cpu_runtime(page_image)
                except Exception as retry_exc:
                    self.last_error = f"{type(retry_exc).__name__}: {retry_exc}"
                    return []
            else:
                self.last_error = f"{type(exc).__name__}: {exc}"
                return []

        parsed = self._parse_result(result)
        if not parsed:
            return []

        blocks: list[dict[str, Any]] = []
        for bbox, payload in parsed:
            text, confidence = payload
            if float(confidence) < min_confidence:
                continue
            blocks.append(
                {
                    "text": text.strip(),
                    "confidence": float(confidence),
                    "bbox": [[int(p[0]), int(p[1])] for p in bbox],
                    "page": page_number,
                }
            )
        return blocks

    def _predict(self, image: np.ndarray) -> Any:
        ocr = self._get_ocr()
        predict = getattr(ocr, "predict", None)
        if callable(predict):
            return predict(image)
        return ocr.ocr(image, cls=self.use_angle_cls)

    @staticmethod
    def _parse_result(result: Any) -> list[tuple[Any, Any]]:
        # PaddleOCR 2.x format: [ [ [bbox, [text, conf]], ... ] ]
        if isinstance(result, list) and result and isinstance(result[0], list):
            first = result[0]
            if first and isinstance(first[0], (list, tuple)) and len(first[0]) == 2:
                return first

        # PaddleOCR 3.x format may return dict-like items.
        if isinstance(result, list) and result and isinstance(result[0], dict):
            normalized: list[tuple[Any, Any]] = []
            for item in result:
                texts = item.get("rec_texts")
                scores = item.get("rec_scores")
                polys = OCREngine._first_present(item, "dt_polys", "rec_polys", "bbox")
                if isinstance(polys, np.ndarray):
                    polys = polys.tolist()

                if isinstance(texts, list):
                    for idx, text in enumerate(texts):
                        if not isinstance(text, str):
                            continue
                        poly = polys[idx] if isinstance(polys, list) and idx < len(polys) else []
                        poly = OCREngine._normalize_poly(poly)
                        score = scores[idx] if isinstance(scores, list) and idx < len(scores) else 0.0
                        normalized.append((poly, [text, float(score)]))
                    continue

                text = item.get("rec_text") or item.get("text") or ""
                score = item.get("rec_score") or item.get("score") or 0.0
                poly = OCREngine._normalize_poly(polys)
                if isinstance(text, str):
                    normalized.append((poly, [text, float(score)]))
            return normalized

        return []

    @staticmethod
    def _first_present(item: dict[str, Any], *keys: str) -> Any:
        for key in keys:
            value = item.get(key)
            if value is not None:
                return value
        return []

    @staticmethod
    def _normalize_poly(poly: Any) -> Any:
        if isinstance(poly, np.ndarray):
            return poly.tolist()
        return poly

    def save_debug_image(
        self,
        input_path: str,
        blocks: list[dict[str, Any]],
        output_path: str,
        preprocess: bool = False,
    ) -> str:
        image = self._load_first_image(input_path)
        if preprocess:
            image = preprocess_image(image)

        for block in blocks:
            pts = np.array(block.get("bbox", []), dtype=np.int32)
            if len(pts) == 4:
                cv2.polylines(image, [pts], True, (0, 255, 0), 2)
                cv2.putText(
                    image,
                    block.get("text", "")[:24],
                    tuple(pts[0]),
                    cv2.FONT_HERSHEY_SIMPLEX,
                    0.5,
                    (0, 0, 255),
                    1,
                    cv2.LINE_AA,
                )

        output = Path(output_path)
        output.parent.mkdir(parents=True, exist_ok=True)
        cv2.imwrite(str(output), image)
        return str(output)

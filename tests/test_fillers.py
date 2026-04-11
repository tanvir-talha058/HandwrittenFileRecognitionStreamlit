from __future__ import annotations

import tempfile
import unittest
from pathlib import Path
from unittest.mock import patch

import fitz
from docx import Document

from form_filler.docx_filler import DOCXFiller
from form_filler.pdf_filler import PDFFiller


class DOCXFillerTests(unittest.TestCase):
    def test_replaces_placeholders_in_runs_tables_and_headers(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            template_path = temp_path / "template.docx"
            output_path = temp_path / "filled.docx"

            document = Document()
            paragraph = document.add_paragraph()
            paragraph.add_run("Applicant: ")
            paragraph.add_run("{{applicant")
            paragraph.add_run("_name}}")

            table = document.add_table(rows=1, cols=1)
            table.rows[0].cells[0].text = "Phone: {{ mobile_number }}"

            header = document.sections[0].header
            header.paragraphs[0].text = "Loan amount: {{loan_amount}}"
            document.save(template_path)

            filler = DOCXFiller(template=str(template_path))
            filler.fill(
                {
                    "applicant_name": "Jane Doe",
                    "mobile_number": "01700000000",
                    "loan_amount": "2500000",
                },
                str(output_path),
            )

            filled = Document(str(output_path))
            body_text = "\n".join(paragraph.text for paragraph in filled.paragraphs)
            table_text = filled.tables[0].rows[0].cells[0].text
            header_text = filled.sections[0].header.paragraphs[0].text

            self.assertIn("Applicant: Jane Doe", body_text)
            self.assertEqual("Phone: 01700000000", table_text)
            self.assertEqual("Loan amount: 2500000", header_text)


class PDFFillerTests(unittest.TestCase):
    def test_ocr_template_anchor_places_text_without_summary_page(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            template_path = temp_path / "blank_template.pdf"
            output_path = temp_path / "filled.pdf"

            document = fitz.open()
            document.new_page(width=595, height=842)
            document.save(str(template_path))
            document.close()

            mocked_lines = [
                {
                    "text": "Applicant Name",
                    "x1": 50.0,
                    "y1": 100.0,
                    "x2": 220.0,
                    "y2": 130.0,
                }
            ]

            with patch.object(PDFFiller, "_extract_template_page_lines", return_value=(mocked_lines, (595, 842))):
                filler = PDFFiller(template=str(template_path))
                filler.fill({"applicant_name": "Jane Doe"}, str(output_path))

            filled = fitz.open(str(output_path))
            try:
                self.assertEqual(1, filled.page_count)
                page_text = filled[0].get_text("text")
                self.assertIn("Jane Doe", page_text)
                self.assertNotIn("Unplaced Extracted Fields", page_text)
            finally:
                filled.close()


if __name__ == "__main__":
    unittest.main()
